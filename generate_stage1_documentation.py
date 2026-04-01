from pathlib import Path
from docx import Document

TEMPLATE_PATH = Path(r"c:/Users/arunt/Documents/Major Project/AgriLink_springboot/2022 major project documentation guidelines _ STAGE_1 FINAL .docx")
OUTPUT_PATH = Path(r"c:/Users/arunt/Documents/Major Project/AgriLink_springboot/AgriLink_Stage1_Project_Documentation.docx")

PROJECT_TITLE = "AGRILINK: A MICROSERVICES-BASED AGRICULTURAL MARKETPLACE PLATFORM"


def replace_text_in_paragraph(paragraph, replacements):
    original = paragraph.text
    updated = original
    for src, dst in replacements.items():
        updated = updated.replace(src, dst)
    if updated != original:
        paragraph.text = updated


def replace_placeholders(doc):
    replacements = {
        "<Title of Project>": PROJECT_TITLE,
        "<Title of the Project>": PROJECT_TITLE,
    }

    for p in doc.paragraphs:
        replace_text_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, replacements)


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def add_paragraphs(doc, items):
    for item in items:
        doc.add_paragraph(item)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(f"- {item}")


def add_numbered(doc, items):
    for idx, item in enumerate(items, start=1):
        doc.add_paragraph(f"{idx}. {item}")


def append_report_content(doc):
    doc.add_page_break()

    add_heading(doc, "ACKNOWLEDGEMENT")
    add_paragraphs(doc, [
        "We express our sincere gratitude to our project guide and the Department of Computer Science and Engineering for their continuous guidance and support throughout the development of this project.",
        "We also thank our institution, faculty members, and peers for providing the resources, feedback, and motivation needed to complete this work.",
        "Finally, we acknowledge all open-source communities and documentation sources that enabled us to design and implement the AgriLink platform effectively.",
    ])

    doc.add_page_break()

    add_heading(doc, "ABSTRACT")
    add_paragraphs(doc, [
        "AgriLink is a full-stack agricultural marketplace platform designed to connect farmers and buyers through a transparent and scalable digital ecosystem. The project adopts a microservices architecture with Spring Boot backend services and a React frontend.",
        "The system provides core capabilities including secure authentication and role management, user profile and verification workflows, farm and field management, produce listing and discovery, cart and checkout, order lifecycle tracking, messaging, notifications, and fraud case management.",
        "The backend is organized as domain-specific services with independent persistence boundaries and Flyway-based schema management, while the frontend provides role-based dashboards and seamless user journeys for customers, farmers, managers, and admins.",
        "This report presents the motivation, problem statement, objectives, requirement analysis, and system design of AgriLink in accordance with Stage-I project documentation guidelines.",
    ])

    doc.add_page_break()

    add_heading(doc, "LIST OF TABLES")
    add_numbered(doc, [
        "Service-to-Port Mapping",
        "Software Requirements",
        "Hardware Requirements",
        "Technology Stack",
        "Core API Module Summary",
    ])

    add_heading(doc, "LIST OF FIGURES")
    add_numbered(doc, [
        "High-Level System Architecture",
        "Authentication and Authorization Flow",
        "Order and Payment Verification Sequence",
        "Inter-Service Communication Overview",
    ])

    add_heading(doc, "ABBREVIATIONS")
    add_bullets(doc, [
        "API - Application Programming Interface",
        "JWT - JSON Web Token",
        "JPA - Java Persistence API",
        "KYC - Know Your Customer",
        "CI/CD - Continuous Integration / Continuous Delivery",
        "DB - Database",
    ])

    add_heading(doc, "SYMBOLS")
    add_paragraphs(doc, [
        "No special mathematical symbols are used in this report.",
    ])

    doc.add_page_break()

    add_heading(doc, "1 INTRODUCTION")

    add_subheading(doc, "1.1 MOTIVATION")
    add_paragraphs(doc, [
        "Conventional agricultural supply chains often include multiple intermediaries, resulting in reduced farmer margins, delayed payments, and limited buyer visibility into produce quality and source.",
        "AgriLink is motivated by the need for a reliable digital marketplace where farmers can directly list produce, buyers can make informed purchases, and transaction operations are traceable across the lifecycle.",
        "The platform is additionally motivated by practical needs such as secure payments, role-based moderation, profile verification, and scalable architecture for future analytics and integrations.",
    ])

    add_subheading(doc, "1.2 PROBLEM STATEMENT")
    add_paragraphs(doc, [
        "The agricultural market lacks an integrated, role-aware platform that supports farmer onboarding, farm data management, listing discovery, transaction processing, and post-order communication under a unified architecture.",
        "Existing ad-hoc systems typically fail in one or more areas: security, maintainability, scalability, fraud handling, and structured data management.",
        "The problem is to design and implement a modular platform that addresses these concerns while remaining deployment-friendly for local and cloud PostgreSQL environments.",
    ])

    add_subheading(doc, "1.3 PROJECT OBJECTIVES")
    add_numbered(doc, [
        "Design a microservices-based backend with clear domain boundaries.",
        "Implement secure user authentication and role-based access control using JWT.",
        "Provide complete user profile and verification workflows for different actor roles.",
        "Enable farm and listing management for farmers.",
        "Implement marketplace search, cart, checkout, and order management flows.",
        "Integrate payment verification and fraud case handling in order workflows.",
        "Provide messaging and notification capabilities for operational communication.",
        "Support both local PostgreSQL and Neon PostgreSQL deployment models.",
    ])

    add_subheading(doc, "1.4 PROJECT REPORT ORGANIZATION")
    add_paragraphs(doc, [
        "Chapter 1 introduces the motivation, problem statement, and objectives.",
        "Chapter 2 presents existing work and limitations relevant to digital agricultural marketplaces.",
        "Chapter 3 captures software, hardware, and user requirements.",
        "Chapter 4 describes system design, architecture, methods, diagrams, and technology stack.",
        "The report concludes with references and appendix material related to implementation artifacts.",
    ])

    doc.add_page_break()

    add_heading(doc, "2 LITERATURE REVIEW")

    add_subheading(doc, "2.1 EXISTING WORK")
    add_paragraphs(doc, [
        "Current e-commerce and agri-trade platforms provide fragmented capabilities such as catalog browsing, basic vendor onboarding, and payment processing. However, many are either marketplace-centric without farm lifecycle support, or farm-centric without robust transactional operations.",
        "Research and industry practices indicate that modular architectures improve maintainability and allow independent scaling of critical services like authentication, order processing, and catalog search.",
        "Modern backend stacks frequently combine Spring Boot, JPA, PostgreSQL, and token-based security due to strong ecosystem support, reliability, and enterprise-grade tooling.",
    ])

    add_subheading(doc, "2.2 LIMITATIONS OF EXISTING WORK")
    add_paragraphs(doc, [
        "Limited role-specific workflows: many solutions do not model separate lifecycle requirements for farmers, buyers, managers, and admins.",
        "Weak traceability and moderation: listing approval, fraud case handling, and structured audit progression are often missing.",
        "Monolithic constraints: tightly coupled systems make it difficult to evolve modules independently or optimize high-load components.",
        "Insufficient communication layer: integrated messaging and notification patterns are frequently absent or externally dependent.",
        "Deployment inflexibility: practical migration paths between local and cloud-hosted PostgreSQL setups are rarely documented.",
    ])

    doc.add_page_break()

    add_heading(doc, "3 REQUIREMENT ANALYSIS")

    add_subheading(doc, "3.1 SOFTWARE REQUIREMENTS")
    software_requirements = [
        "Operating System: Windows 10/11, Linux, or macOS",
        "Backend Runtime: Java (project modules configured with Spring Boot ecosystem)",
        "Build Tool: Maven (multi-module)",
        "Frontend Runtime: Node.js and npm",
        "Frontend Framework: React",
        "Database: PostgreSQL (local/docker/neon)",
        "Containerization: Docker and Docker Compose",
        "Version Control: Git",
        "Testing: JUnit/Mockito and API-level validations",
    ]
    add_bullets(doc, software_requirements)

    add_subheading(doc, "3.2 HARDWARE REQUIREMENTS")
    hardware_requirements = [
        "Processor: Modern multi-core CPU (Intel i5/Ryzen 5 equivalent or above)",
        "RAM: Minimum 8 GB (16 GB recommended for running multiple services)",
        "Storage: Minimum 20 GB free space (SSD recommended)",
        "Network: Stable internet connection for dependency downloads and cloud DB access",
    ]
    add_bullets(doc, hardware_requirements)

    add_subheading(doc, "3.3 USER REQUIREMENTS")
    add_paragraphs(doc, [
        "The platform supports the following user personas:",
    ])
    add_bullets(doc, [
        "Farmer: manage profile, farms, fields, crop plans, and listings; track sales.",
        "Buyer/Customer: browse listings, manage cart, checkout, and track purchases.",
        "Manager/Admin: verify profiles, moderate users/listings, monitor fraud cases.",
        "System Operator: deploy services, configure environment variables, monitor logs.",
    ])
    add_paragraphs(doc, [
        "Key functional requirements include secure login, role-based authorization, listing search, order state transitions, payment verification, messaging, and notification delivery.",
        "Key non-functional requirements include reliability, maintainability, scalability, consistency, and secure handling of credentials and tokens.",
    ])

    doc.add_page_break()

    add_heading(doc, "4 SYSTEM DESIGN")

    add_subheading(doc, "4.0 PROPOSED SYSTEM ARCHITECTURE")
    add_paragraphs(doc, [
        "AgriLink follows a microservices architecture where each domain capability is implemented as an independent Spring Boot service with dedicated data ownership and migration strategy.",
        "The frontend communicates with backend REST APIs, and backend services interact via configured service URLs where required. Docker compose files support local and Neon-based setups.",
        "Core backend services and default ports are shown below:",
    ])

    table = doc.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "Service"
    header[1].text = "Default Port"
    header[2].text = "Primary Responsibility"

    rows = [
        ("auth-service", "8081", "Authentication and JWT authorization"),
        ("user-service", "8082", "Profiles, KYC, and role-specific user data"),
        ("farm-service", "8083", "Farm, field, crop plan, and farm analytics flows"),
        ("marketplace-service", "8084", "Listings, categories, reviews, wishlist, demand insights"),
        ("order-service", "8085", "Cart, checkout, payments, order lifecycle, fraud cases"),
        ("notification-service", "8087", "Notifications, templates, messaging, email integration"),
        ("frontend", "3000", "Role-based user interface and dashboard workflows"),
    ]

    for svc, port, role in rows:
        cells = table.add_row().cells
        cells[0].text = svc
        cells[1].text = port
        cells[2].text = role

    add_subheading(doc, "4.1 PROPOSED METHODS / ALGORITHMS")
    add_numbered(doc, [
        "Authentication method: JWT token issuance and validation for secure stateless sessions.",
        "Role-based authorization method: endpoint-level access checks for FARMER, BUYER/CUSTOMER, MANAGER, ADMIN.",
        "Listing retrieval method: searchable filtering with controlled sort fields and paginated responses.",
        "Checkout and payment method: checkout initialization, gateway transaction reference binding, signature verification before completion.",
        "Fraud management method: case creation, status progression, and investigation note logging for admin moderation.",
        "Notification method: event-driven user alert creation and message thread management.",
    ])

    add_subheading(doc, "4.2 CLASS / USE CASE / ACTIVITY / SEQUENCE DIAGRAMS")
    add_paragraphs(doc, [
        "Stage-I design artifacts are summarized below in textual form:",
        "Use Case Scope: Farmer listing management, buyer ordering, admin moderation, and support messaging.",
        "Sequence Example: Buyer checkout flow traverses cart validation -> checkout initialization -> payment gateway callback verification -> order state update -> notification dispatch.",
        "Class-Level Domains: Auth entities, user profile entities, farm entities, listing entities, order/payment entities, notification/messaging entities.",
        "Activity Overview: User onboarding, listing approval, order lifecycle transitions, and fraud-case closure loops.",
        "Note: Visual UML diagrams and sequence diagrams can be inserted in this section as final design figures.",
    ])

    add_subheading(doc, "4.3 DATASETS AND TECHNOLOGY STACK")
    add_paragraphs(doc, [
        "Dataset and seed sources used in development:",
    ])
    add_bullets(doc, [
        "seed-data.sql",
        "seed-data-fixed.sql",
        "seed-data-neon.sql",
        "seed-agrilink-products.sql",
        "insert-admin.sql",
    ])

    add_paragraphs(doc, [
        "Technology stack summary:",
    ])

    tech_table = doc.add_table(rows=1, cols=2)
    tech_header = tech_table.rows[0].cells
    tech_header[0].text = "Layer"
    tech_header[1].text = "Technologies"

    tech_rows = [
        ("Frontend", "React, React Router, Axios, Context API, React Toastify"),
        ("Backend", "Spring Boot, Spring Security, Spring Data JPA, Flyway"),
        ("Database", "PostgreSQL (local/docker), Neon PostgreSQL"),
        ("Build and Dependency", "Maven (multi-module), npm"),
        ("Containerization", "Docker, Docker Compose"),
        ("Testing", "JUnit, Mockito, service-level tests"),
    ]
    for layer, stack in tech_rows:
        row = tech_table.add_row().cells
        row[0].text = layer
        row[1].text = stack

    doc.add_page_break()

    add_heading(doc, "REFERENCES")
    add_numbered(doc, [
        "AgriLink repository README and module documentation.",
        "Spring Boot Reference Documentation.",
        "Spring Security Reference Documentation.",
        "React Official Documentation.",
        "PostgreSQL Official Documentation.",
        "Flyway Migration Documentation.",
        "Docker and Docker Compose Documentation.",
    ])

    doc.add_page_break()

    add_heading(doc, "APPENDIX: IMPLEMENTATION NOTES")
    add_paragraphs(doc, [
        "A. Project Structure: Multi-module backend with domain-oriented microservices and one React frontend application.",
        "B. Environment Profiles: Local profile and Neon profile support for cloud-hosted PostgreSQL migration.",
        "C. Integration Notes: Marketplace and order flows depend on secure inter-service communication and consistent DTO contracts.",
        "D. Operational Scripts: The repository includes scripts for test execution and Neon-based service startup orchestration.",
        "E. Scope Note: IoT references exist in documentation and build artifacts; active source participation can be validated against current workspace state before Stage-II expansion.",
    ])


def main():
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    replace_placeholders(doc)
    append_report_content(doc)
    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
